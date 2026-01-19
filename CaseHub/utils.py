import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import io

def generate_case_study(api_key, industry, topic, difficulty, company_size):
    """
    Generates a structured case study using Google Gemini.
    """
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')

        prompt = f"""
        You are an expert case study writer for academic and professional training purposes.
        Create a detailed, realistic business case study based on the following parameters:
        
        - Industry: {industry}
        - Specific Topic: {topic}
        - Difficulty Level: {difficulty}
        - Company Size: {company_size}
        
        The case study MUST strictly follow this structure:
        
        # [Case Study Title]
        
        ## Executive Summary
        (A brief overview of the case)
        
        ## Background
        (Context about the company, market position, and history)
        
        ## The Challenge
        (Detailed description of the problem/crisis/opportunity)
        
        ## Proposed Solution
        (Strategic options or the implemented solution to be analyzed)
        
        ## Discussion Questions
        (3-5 distinct questions for students/professionals to answer)
        
        Make the content engaging, professional, and suitable for the specified difficulty level.
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error creating case study: {str(e)}"

def create_pdf(markdown_content):
    """
    Converts markdown content to a simple PDF.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    
    # Simple markdown-like processing (very basic)
    # Replace common unicode characters that fpdf (latin-1) struggles with
    replacements = {
        '\u2013': '-',  # en dash
        '\u2014': '--', # em dash
        '\u2018': "'",  # left single quote
        '\u2019': "'",  # right single quote
        '\u201c': '"',  # left double quote
        '\u201d': '"',  # right double quote
        '\u2022': '*',  # bullet
        '\u2026': '...' # ellipsis
    }
    
    for line in markdown_content.split('\n'):
        # Sanitize line
        for char, replacement in replacements.items():
            line = line.replace(char, replacement)
        
        # Ensure it's latin-1 compatible, replacing unknown chars with ?
        line = line.encode('latin-1', 'replace').decode('latin-1')
        
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
            
        if line.startswith('# '):
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, line[2:], ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('## '):
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, line[3:], ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('### '):
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, line[4:], ln=True)
            pdf.set_font("Arial", size=12)
        else:
            # Handle standard text wrap
            pdf.multi_cell(0, 10, line)
            
    return pdf.output(dest='S').encode('latin-1') # Return bytes

def create_docx(markdown_content):
    """
    Converts markdown content to a Word document.
    """
    doc = Document()
    
    for line in markdown_content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        else:
            doc.add_paragraph(line)
            
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
